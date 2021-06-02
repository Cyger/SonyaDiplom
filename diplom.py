#!/usr/bin/python3
from docx import Document
import pandas
import xlrd

def doc():
    document = Document('test_doc.docx')

    line = 'Объем дисциплины '
    new_line = 'Объем дисциплины '

    for paragraph in document.paragraphs:
        #if line in paragraph.text:
        print(paragraph.text)
            #paragraph.text = new_line
            #print(paragraph.text)

    #document.save('test.docx')
# 4 8 13 17 22 26 31 35
#openpyxl
def xls():
    file = pandas.ExcelFile('IST4.xls')

    courses_index = [4,8,13,17,22,26,31,35]
    subj = "Технологии программирования"
    dr = pandas.read_excel(file, 'Диаграмма курсов')
   
    attestation = []

    for i in range(8):
        for j in range(350):
            volume = str(dr.iat[j,courses_index[i]])
            if subj in volume:
                #if ('За' in volume) and ('ЗаО' not in volume):
                    #attestation.append('Сем ' + str(i+1) + ' - За')
                    #break
                #if 'ЗаО' in volume:
                    #attestation.append('Сем ' + str(i+1) + ' - ЗаО')
                    #pring(volume)
                if 'За' in volume and 'ЗаО' not in volume:
                    attestation.append('Сем ' + str(i+1) + ' - За')
                    break
                if 'ЗаО' in volume:
                    attestation.append('Сем ' + str(i+1) + ' - ЗаО')
                    break
                if 'Экз' in volume:
                    attestation.append('Сем ' + str(i+1) + ' - Экз')
                    break


    
def discipline_volume(discipline):

    education_level = "Бакалавриат"
    course_discipline = "Информатика"
    course_code = "09.03.02"
    course = "Информационные системы и технологии"
    education_start_year = "2018"
    graduate_department = "ГИС"
    developer_department = "ГИС"
    programm_variant = "5"

   
    xls_file = pandas.ExcelFile('IST4.xls')
    df = pandas.read_excel(xls_file, 'План')
    doc_file = Document('test_doc.docx')

###   Рабочая программа дисциплини
    search_line = 'РАБОЧАЯ ПРОГРАММА ДИСЦИПЛИНЫ'
    paragraph_index = 0
    for paragraph in doc_file.paragraphs:
        if paragraph_index == 1:
            paragraph.text += '   ' + "РПД1"
            break
        if search_line in paragraph.text:
            paragraph_index = 1

###   Рабочая программа дисциплини 2
    search_line = 'РАБОЧАЯ ПРОГРАММА ДИСЦИПЛИНЫ'
    paragraph_index = 0
    for paragraph in doc_file.paragraphs:
        if paragraph_index == 3:
            paragraph.text += '   ' + "РПД2"
            break
        if paragraph_index == 2:
            paragraph_index += 1
        if search_line in paragraph.text:
            paragraph_index += 1

###   База

    search_line = 'Направление подготовки'
    for paragraph in doc_file.paragraphs:
        if search_line in paragraph.text:
            paragraph.text += '   ' + course_code + ' ' + course
            break

    search_line = 'Направленность: '
    for paragraph in doc_file.paragraphs:
        if search_line in paragraph.text:
            paragraph.text += '   ' + "НАПРАВЛЕННОСТЬ"
            break

    search_line = 'Год начала подготовки '
    for paragraph in doc_file.paragraphs:
        if search_line in paragraph.text:
            paragraph.text += '   ' + education_start_year
            break

    search_line = 'Выпускающая кафедра'
    for paragraph in doc_file.paragraphs:
        if search_line in paragraph.text:
            paragraph.text += '   ' + graduate_department
            break

    search_line = 'Кафедра-разработчик'
    for paragraph in doc_file.paragraphs:
        if search_line in paragraph.text:
            paragraph.text += '   ' + developer_department
            break


###   Промежуточная аттестация
    courses_index = [4,8,13,17,22,26,31,35]
    dr = pandas.read_excel(xls_file, 'Диаграмма курсов')
    attestation = []

    for i in range(8):
        for j in range(350):
            volume = str(dr.iat[j,courses_index[i]])
            if discipline in volume:
                #if ('За' in volume) and ('ЗаО' not in volume):
                    #attestation.append('Сем ' + str(i+1) + ' - За')
                    #break
                #if 'ЗаО' in volume:
                    #attestation.append('Сем ' + str(i+1) + ' - ЗаО')
                    #pring(volume)
                if 'За' in volume and 'ЗаО' not in volume:
                    attestation.append('Сем ' + str(i+1) + ' - За')
                    break
                if 'ЗаО' in volume:
                    attestation.append('Сем ' + str(i+1) + ' - ЗаО')
                    break
                if 'Экз' in volume:
                    attestation.append('Сем ' + str(i+1) + ' - Экз')
                    break

    search_line = 'Промежуточная аттестация'
    attestation_string = ''
    for paragraph in doc_file.paragraphs:
        if search_line in paragraph.text:
            for i in attestation:
                attestation_string += '   ' + i
            paragraph.text += '   ' + attestation_string
            break



###   Объем дисциплины
    for i in range(df.shape[0]):
        if df.iat[i,3] == discipline:
            discipline_volume_amount = df.iat[i,14]

    search_line = 'Объем дисциплины '
    #new_line = f"jksjkdjf {volume}"
    for paragraph in doc_file.paragraphs:
        if search_line in paragraph.text:
            paragraph.text += discipline_volume_amount



    doc_file.save('test_output.docx')


discipline_volume("Технологии программирования")
#xls()

